import numpy as np
import pandas as pd
from docx import Document

data = pd.read_excel("Qualitative data analysis.xlsx", sheet_name='currentValue')
data.columns.values[2:7] = ['19Feb','20Feb', '21Feb', '22Feb', '23Feb']
print(data.columns)

# data.rename(columns={'currentValue': 'Feb-19', 'currentValue.1': 'Feb-20', 'currentValue.2': 'Feb-21'}, inplace=True)
currentValue_data = data[['scripCode', 'companyName', '19Feb','20Feb', '21Feb', '22Feb', '23Feb']]

highest_values = []
highest_column_names = []
lowest_values = []
lowest_column_names = []

for index, row in currentValue_data.iterrows():
    highest_values.append(row[['19Feb','20Feb', '21Feb', '22Feb', '23Feb']].max())
    highest_column_names.append(row[['19Feb','20Feb', '21Feb', '22Feb', '23Feb']].idxmax())
    lowest_values.append(row[['19Feb','20Feb', '21Feb', '22Feb', '23Feb']].min())
    lowest_column_names.append(row[['19Feb','20Feb', '21Feb', '22Feb', '23Feb']].idxmin())

#doc = Document()

#doc.add_heading('Current Value Analysis', level=1)

# for i in range(len(currentValue_data)):
#     doc.add_paragraph(f"Scrip Code: {currentValue_data['scripCode'][i]}")
#     doc.add_paragraph(f"Company Name: {currentValue_data['companyName'][i]}")
#     doc.add_paragraph(f"Highest Value: {highest_values[i]}, Column Name: {highest_column_names[i]}")
#     doc.add_paragraph(f"Lowest Value: {lowest_values[i]}, Column Name: {lowest_column_names[i]}")
#     doc.add_paragraph("")

# doc.save('currentValue_analysis.docx')
currentValue_data['Highest Value'] = highest_values
currentValue_data['Highest Value Column'] = highest_column_names
currentValue_data['Lowest Value'] = lowest_values
currentValue_data['Lowest Value Column'] = lowest_column_names

currentValue_data.to_excel('currentValue_Analysis.xlsx', sheet_name='currentValue', index=True)

